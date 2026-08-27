#!/usr/bin/env python3
"""
Biography Writer Skill - Main Script
Transforms biographical data into a professional Word document with chapters, photos, TOC, and timeline.
"""

import json
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class BiographyGenerator:
    """Generate professional biographical books from structured interview data."""
    
    def __init__(self, person_name: str):
        self.person_name = person_name
        self.data = {
            'nama': person_name,
            'masa_kecil': {},
            'pendidikan': {},
            'karir_awal': {},
            'pencapaian_utama': {},
            'kehidupan_keluarga': {},
            'warisan': {}
        }
        self.chapters = []
        self.timeline_events = []
        self.photos = {}  # chapter_name -> list of photo paths
        
        # Default chapter structure
        self.chapter_order = [
            ('masa_kecil', 'MASA KECIL'),
            ('pendidikan', 'PENDIDIKAN'),
            ('karir_awal', 'KARIR AWAL'),
            ('pencapaian_utama', 'PENCAPAIAN UTAMA'),
            ('kehidupan_keluarga', 'KEHIDUPAN KELUARGA'),
            ('warisan', 'WARISAN DAN LEGASI')
        ]
    
    def add_chapter_data(self, chapter_key: str, data: Dict) -> None:
        """Add data for a specific chapter."""
        if chapter_key in self.data:
            self.data[chapter_key] = data
    
    def add_photo(self, chapter_key: str, photo_path: str, caption: str = "") -> None:
        """Add a photo to a chapter."""
        if chapter_key not in self.photos:
            self.photos[chapter_key] = []
        self.photos[chapter_key].append({
            'path': photo_path,
            'caption': caption
        })
    
    def add_timeline_event(self, year: int, event: str, month: Optional[int] = None) -> None:
        """Add an event to the timeline."""
        self.timeline_events.append({
            'year': year,
            'month': month,
            'event': event
        })
        # Sort by year and month
        self.timeline_events.sort(key=lambda x: (x['year'], x['month'] or 0))
    
    def generate_chapter_narrative(self, chapter_data: Dict, chapter_name: str) -> str:
        """
        Generate semi-formal narrative paragraphs from chapter data.
        This should be called with answers from the interview.
        """
        # This would typically be done by Claude using the interview data
        # For now, returning placeholder that would be replaced by actual narrative
        return f"Chapter narrative for {chapter_name} to be generated from data."
    
    def create_document(self) -> Document:
        """Create and format the biography document."""
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # 1. Title Page
        self._add_title_page(doc)
        
        # 2. Table of Contents
        self._add_table_of_contents(doc)
        
        # 3. Timeline/Sidebar (if events exist)
        if self.timeline_events:
            self._add_timeline_section(doc)
        
        # 4. Chapters
        for chapter_key, chapter_title in self.chapter_order:
            if chapter_key in self.data and self.data[chapter_key]:
                self._add_chapter(doc, chapter_key, chapter_title)
        
        return doc
    
    def _add_title_page(self, doc: Document) -> None:
        """Add professional title page."""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(self.person_name)
        title_run.font.size = Pt(48)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('Biografi Lengkap')
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = RGBColor(127, 140, 141)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Date generated
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f'Digenerate: {datetime.now().strftime("%d %B %Y")}')
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(149, 165, 166)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document) -> None:
        """Add Table of Contents."""
        toc_title = doc.add_paragraph()
        toc_title_run = toc_title.add_run('DAFTAR ISI')
        toc_title_run.font.size = Pt(14)
        toc_title_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add TOC entries
        for chapter_key, chapter_title in self.chapter_order:
            if chapter_key in self.data and self.data[chapter_key]:
                toc_entry = doc.add_paragraph(
                    f'{chapter_title}',
                    style='List Number'
                )
                toc_entry.paragraph_format.left_indent = Inches(0.25)
        
        if self.timeline_events:
            doc.add_paragraph('Timeline Kehidupan', style='List Number')
        
        doc.add_page_break()
    
    def _add_timeline_section(self, doc: Document) -> None:
        """Add timeline/sidebar section."""
        title = doc.add_paragraph()
        title_run = title.add_run('TIMELINE KEHIDUPAN')
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        
        doc.add_paragraph()
        
        # Add timeline events
        for event in self.timeline_events:
            date_str = f"{event['year']}"
            if event['month']:
                months = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
                         'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
                date_str = f"{months[event['month']]} {event['year']}"
            
            # Timeline entry
            entry = doc.add_paragraph()
            entry_run = entry.add_run(f"• {date_str} - {event['event']}")
            entry_run.font.size = Pt(11)
            entry.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_page_break()
    
    def _add_chapter(self, doc: Document, chapter_key: str, chapter_title: str) -> None:
        """Add a chapter with narrative, photos, and formatting."""
        # Chapter title
        title = doc.add_paragraph()
        title_run = title.add_run(chapter_title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title.paragraph_format.space_before = Pt(12)
        title.paragraph_format.space_after = Pt(12)
        
        # Chapter narrative (from self.data)
        chapter_data = self.data[chapter_key]
        if isinstance(chapter_data, dict) and 'narasi' in chapter_data:
            narasi = chapter_data['narasi']
            # Split into paragraphs if it's a long text
            for para_text in narasi.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p_format = p.paragraph_format
                    p_format.line_spacing = 1.5
                    p_format.space_after = Pt(12)
        
        # Add photos if they exist
        if chapter_key in self.photos:
            for photo_info in self.photos[chapter_key]:
                try:
                    # Add photo
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(photo_info['path'], width=Inches(4))
                    
                    # Add caption
                    if photo_info['caption']:
                        caption = doc.add_paragraph(photo_info['caption'])
                        caption_format = caption.paragraph_format
                        caption_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.runs[0]
                        caption_run.font.size = Pt(9)
                        caption_run.font.italic = True
                        caption_run.font.color.rgb = RGBColor(127, 140, 141)
                except Exception as e:
                    print(f"Warning: Could not add photo for {chapter_key}: {e}")
                    doc.add_paragraph(f"[Foto: {photo_info['caption']}]")
        
        doc.add_page_break()
    
    def save(self, filename: str) -> str:
        """Generate and save the document."""
        doc = self.create_document()
        doc.save(filename)
        return filename


def main():
    """Example usage of BiographyGenerator."""
    # Create generator
    bg = BiographyGenerator("John Doe")
    
    # Add sample data for each chapter
    bg.add_chapter_data('masa_kecil', {
        'narasi': '''John Doe dilahirkan pada 15 Januari 1975 di Jakarta. Orang tuanya, 
Bapak Sutrisno dan Ibu Endang, menyambut kedatangannya dengan kebahagiaan yang mendalam. 
Masa kecilnya diisi dengan kenangan indah di tengah keluarga yang hangat dan penuh kasih sayang.

Sebagai anak tertua, John sejak dini menunjukkan tanggung jawab yang luar biasa. 
Dia sangat dekat dengan ayahnya yang bekerja di industri minyak, dan sering mendengarkan 
kisah-kisah tentang dunia kerja dan kehidupan. Ibunya, seorang pendidik, membesarkannya 
dengan nilai-nilai pembelajaran, kesabaran, dan empati terhadap sesama.'''
    })
    
    bg.add_chapter_data('pendidikan', {
        'narasi': '''Perjalanan pendidikan John dimulai dari SD Negeri 1 Jakarta pada tahun 1981. 
Sejak awal, ia menunjukkan prestasi akademik yang luar biasa, selalu berada di peringkat teratas kelasnya. 
Guru-gurunya mengenali potensi briliannya dan memberikan perhatian khusus untuk mengembangkan bakat-bakat tersembunyi.

Di SMP Santo Louis (1987-1990), John mulai menunjukkan ketertarikan yang mendalam pada sains dan teknologi. 
Dia aktif dalam kompetisi akademik tingkat kota dan bahkan nasional. Selama masa SMP, John juga belajar 
untuk berinteraksi dengan berbagai karakter teman sebaya, keterampilan yang akan berguna di masa depannya.'''
    })
    
    # Add timeline events
    bg.add_timeline_event(1975, 'Lahir di Jakarta', month=1)
    bg.add_timeline_event(1981, 'Memulai SD Negeri 1 Jakarta')
    bg.add_timeline_event(1987, 'Melanjutkan ke SMP Santo Louis')
    bg.add_timeline_event(1990, 'Lulus SMP dengan prestasi terbaik')
    
    # Generate document
    output_file = 'biography_john_doe.docx'
    result = bg.save(output_file)
    print(f"Biography generated successfully: {result}")


if __name__ == '__main__':
    main()
